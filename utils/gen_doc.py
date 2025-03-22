from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml
import io
import tempfile

def export_to_word(response_content, subdomain_definition, science_goal, context, max_tokens, temperature, top_p, frequency_penalty, presence_penalty):
    doc = Document()
    
    # Add a title (optional, you can remove this if not needed)
    doc.add_heading('AI Generated SCDD', 0)

    # Insert the Subdomain Definition at the top
    doc.add_heading('Subdomain Definition:', level=1)
    doc.add_paragraph(subdomain_definition)

    # Insert the Science Goal at the top
    doc.add_heading('Science Goal:', level=1)
    doc.add_paragraph(science_goal)

    # Insert the User-defined Context
    doc.add_heading('User-defined Context:', level=1)
    doc.add_paragraph(context)

    # Insert Model Parameters
    doc.add_heading('Model Parameters:', level=1)
    doc.add_paragraph(f"Max Tokens: {max_tokens}")
    doc.add_paragraph(f"Temperature: {temperature}")
    doc.add_paragraph(f"Top-p: {top_p}")
    doc.add_paragraph(f"Frequency Penalty: {frequency_penalty}")
    doc.add_paragraph(f"Presence Penalty: {presence_penalty}")

    # Split the response into sections based on ### headings
    sections = response_content.split('### ')
    
    for section in sections:
        if section.strip():
            # Handle the "Observations Requirements Table" separately with proper formatting
            if 'Observations Requirements Table' in section:
                doc.add_heading('Observations Requirements Table', level=1)
                
                # Extract table lines
                table_lines = section.split('\n')[2:]  # Start after the heading line
                
                # Check if it's an actual table (split lines by '|' symbol)
                table_data = [line.split('|')[1:-1] for line in table_lines if '|' in line]
                
                if table_data:
                    # Add table to the document
                    table = doc.add_table(rows=len(table_data), cols=len(table_data[0]))
                    table.style = 'Table Grid'
                    for i, row in enumerate(table_data):
                        for j, cell_text in enumerate(row):
                            cell = table.cell(i, j)
                            cell.text = cell_text.strip()
                            # Apply text wrapping for each cell
                            cell._element.get_or_add_tcPr().append(parse_xml(r'<w:tcW w:w="2500" w:type="pct" ' + nsdecls('w') + '/>'))
                
                # Process any paragraphs that follow the table
                paragraph_after_table = '\n'.join([line for line in table_lines if '|' not in line and line.strip()])
                if paragraph_after_table:
                    doc.add_paragraph(paragraph_after_table.strip())
            
            # Handle the "ADS References" section
            elif section.startswith('ADS References'):
                doc.add_heading('ADS References', level=1)
                references = section.split('\n')[1:]  # Skip the heading
                for reference in references:
                    if reference.strip():
                        doc.add_paragraph(reference.strip())
            
            # Add all other sections as plain paragraphs
            else:
                doc.add_paragraph(section.strip())
    
    # Save the document to a temporary file
    temp_file = tempfile.NamedTemporaryFile(delete=False, suffix=".docx")
    doc.save(temp_file.name)
    
    return temp_file.name