
import gradio as gr
from transformers import AutoTokenizer, AutoModel
from openai import OpenAI
import os
import numpy as np
from sklearn.metrics.pairwise import cosine_similarity
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml
import io
import tempfile
from astroquery.nasa_ads import ADS
import pyvo as vo
import pandas as pd
from pinecone import Pinecone
import logging
import re

from langchain_openai import ChatOpenAI
from langchain_openai import OpenAIEmbeddings
llm = ChatOpenAI(model="gpt-4o")
embeddings = OpenAIEmbeddings()
from ragas import EvaluationDataset
from ragas import evaluate
from ragas.llms import LangchainLLMWrapper
evaluator_llm = LangchainLLMWrapper(llm)
from ragas.metrics import LLMContextRecall, ContextRelevance, Faithfulness, ResponseRelevancy, FactualCorrectness

# Load the NASA-specific bi-encoder model and tokenizer
bi_encoder_model_name = "nasa-impact/nasa-smd-ibm-st-v2"
bi_tokenizer = AutoTokenizer.from_pretrained(bi_encoder_model_name)
bi_model = AutoModel.from_pretrained(bi_encoder_model_name)

# Set up OpenAI client
api_key = os.getenv('OPENAI_API_KEY')
client = OpenAI(api_key=api_key)

# Set up NASA ADS token
ADS.TOKEN = os.getenv('ADS_API_KEY')  # Ensure your ADS API key is stored in environment variables

# Pinecone setup
pinecone_api_key = os.getenv('PINECONE_API_KEY')
pc = Pinecone(api_key=pinecone_api_key)
index_name = "scdd-index"
index = pc.Index(index_name)

# Define system message with instructions
system_message = """
You are ExosAI, an advanced assistant specializing in Exoplanet and Astrophysics research.

Generate a **detailed and structured** response based on the given **retrieved context and user input**, incorporating key **observables, physical parameters, and technical requirements**. Organize the response into the following sections:

1. **Science Objectives**: Define key scientific objectives related to the science context and user input.
2. **Physical Parameters**: Outline the relevant physical parameters (e.g., mass, temperature, composition).
3. **Observables**: Specify the key observables required to study the science context.
4. **Description of Desired Observations**: Detail the observational techniques, instruments, or approaches necessary to gather relevant data.
5. **Observations Requirements Table**: Generate a table relevant to the Science Objectives, Physical Parameters, Observables and Description of Desired Observations with the following columns and at least 7 rows:
    - Wavelength Band: Should only be UV, Visible and Infrared).
    - Instrument: Should only be Imager, Spectrograph, Polarimeter and Coronagraph).
    - Necessary Values: The necessary values or parameters (wavelength range, spectral resolution where applicable, spatial resolution where applicable, contrast ratio where applicable).
    - Desired Values: The desired values or parameters (wavelength range, spectral resolution where applicable, spatial resolution where applicable).
    - Number of Objects Observed: Estimate the number of objects that need to be observed for a statistically meaningful result or for fulfilling the science objective.
    - Justification: Detailed scientific explanation of why these observations are important for the science objectives.
    - Comments: Additional notes or remarks regarding each observation.

#### **Table Format** 

| Wavelength Band      | Instrument                         | Necessary Values                   | Desired Values                  | Number of Objects Observed      | Justification     | Comments |
|----------------------|------------------------------------|------------------------------------|---------------------------------|---------------------------------|-------------------|----------|

#### **Guiding Constraints (Exclusions & Prioritization)**
- **Wavelength Band Restriction:** Only include **UV, Visible, and Infrared** bands.
- **Instrument Restriction:** Only include **Imager, Spectrograph, Polarimeter, and Coronagraph**.
- **Wavelength Limits:** Prioritize wavelengths between **100 nanometers (nm) and 3 micrometers (μm)**.
- **Allowed Instruments:** **Only include** observations from **direct imaging, spectroscopy, and polarimetry.** **Exclude** transit and radial velocity methods.
- **Exclusion of Existing Facilities:** **Do not reference** existing observatories such as JWST, Hubble, or ground-based telescopes. This work pertains to a **new mission**.
- **Spectral Resolution Constraint:** Limit spectral resolution (**R**) to the range **10,000 – 50,000**.
- **Contrast Ratio:** Limit contrast ratio to the range **10^4 - 10^6**.
- **Estimate the "Number of Objects Observed" based on the observational strategy, parameters, instruments, statistical requirements, and feasibility.**
- **Ensure that all parameters remain scientifically consistent.**
- **Include inline references wherever available**. Especially in the Justification column.
- **Pay attention to the retrieved context**.

**Use this table format as a guideline, generate a detailed table dynamically based on the input.**. Ensure that all values align with the provided constraints and instructions.

**Include inline references wherever available**. Especially in the Justification column.

Ensure the response is **structured, clear, and observation requirements table follows this format**. **All included parameters must be scientifically consistent with each other.**
"""

# Function to encode query text
def encode_query(text):
    inputs = bi_tokenizer(text, return_tensors='pt', padding=True, truncation=True, max_length=128)
    outputs = bi_model(**inputs)
    embedding = outputs.last_hidden_state.mean(dim=1).detach().numpy().flatten()
    embedding /= np.linalg.norm(embedding)
    return embedding.tolist()


# Context retrieval function using Pinecone
def retrieve_relevant_context(user_input, context_text, science_objectives="", top_k=3):
    query_text = f"Science Goal: {user_input}\nContext: {context_text}\nScience Objectives: {science_objectives}" if science_objectives else f"Science Goal: {user_input}\nContext: {context_text}"
    query_embedding = encode_query(query_text)

    # Pinecone query
    query_response = index.query(
        vector=query_embedding,
        top_k=top_k,
        include_metadata=True
    )

    retrieved_context = "\n\n".join([match['metadata']['text'] for match in query_response.matches])

    if not retrieved_context.strip():
        return "No relevant context found for the query."

    return retrieved_context

def clean_retrieved_context(raw_context):
    # Remove unnecessary line breaks within paragraphs
    cleaned = raw_context.replace("-\n", "").replace("\n", " ")

    # Remove extra spaces clearly
    cleaned = re.sub(r'\s+', ' ', cleaned)

    # Return explicitly cleaned context
    return cleaned.strip()


def extract_keywords_with_gpt(context, max_tokens=100, temperature=0.3):
    
    keyword_prompt = f"Extract 3 most important scientific keywords from the following user query:\n\n{context}"
    
    response = client.chat.completions.create(
        model="gpt-4",
        messages=[
            {"role": "system", "content": "You are an expert in identifying key scientific terms and concepts."},
            {"role": "user", "content": keyword_prompt}
        ],
        max_tokens=max_tokens,
        temperature=temperature
    )
    
    extracted_keywords = response.choices[0].message.content.strip()
    
    cleaned_keywords = re.sub(r'\d+\.\s*', '', extracted_keywords)

    keywords_list = [kw.strip() for kw in cleaned_keywords.split("\n") if kw.strip()]
    
    return keywords_list

def fetch_nasa_ads_references(ads_query):
    """Fetch relevant NASA ADS papers and format them for readability."""
    try:
        # Query NASA ADS for relevant papers
        papers = ADS.query_simple(ads_query)

        if not papers or len(papers) == 0:
            return [("No results found", "N/A", "N/A", "N/A", "N/A", "N/A")]

        # Include authors in the references
        references = []
        for paper in papers[:5]:  # Limit to 5 references
            title = paper.get('title', ['Title not available'])[0]
            abstract = paper.get('abstract', 'Abstract not available')
            authors = ", ".join(paper.get('author', [])[:3]) + (" et al." if len(paper.get('author', [])) > 3 else "")
            bibcode = paper.get('bibcode', 'N/A')
            pub = paper.get('pub', 'Unknown Journal')
            pubdate = paper.get('pubdate', 'Unknown Date')

            references.append((title, abstract, authors, bibcode, pub, pubdate))

        return references

    except Exception as e:
        logging.error(f"Error fetching ADS references: {str(e)}")
        return [("Error fetching references", "See logs for details", "N/A", "N/A", "N/A", "N/A")]

def fetch_exoplanet_data():
    # Connect to NASA Exoplanet Archive TAP Service
    tap_service = vo.dal.TAPService("https://exoplanetarchive.ipac.caltech.edu/TAP")

    # Query to fetch all columns from the pscomppars table
    ex_query = """
        SELECT TOP 10 pl_name, hostname, sy_snum, sy_pnum, discoverymethod, disc_year, disc_facility, pl_controv_flag, pl_orbper, pl_orbsmax, pl_rade, pl_bmasse, pl_orbeccen, pl_eqt, st_spectype, st_teff, st_rad, st_mass, ra, dec, sy_vmag
        FROM pscomppars
    """
    # Execute the query
    qresult = tap_service.search(ex_query)

    # Convert to a Pandas DataFrame
    ptable = qresult.to_table()
    exoplanet_data = ptable.to_pandas()

    return exoplanet_data

def generate_response(user_input, science_objectives="", relevant_context="", references=[], max_tokens=150, temperature=0.7, top_p=0.9, frequency_penalty=0.5, presence_penalty=0.0):
    # Case 1: Both relevant context and science objectives are provided
    if relevant_context and science_objectives.strip():
        combined_input = f"Scientific Context: {relevant_context}\nUser Input: {user_input}\nScience Objectives (User Provided): {science_objectives}\n\nPlease generate only the remaining sections as per the defined format."
    
    # Case 2: Only relevant context is provided
    elif relevant_context:
        combined_input = f"Scientific Context: {relevant_context}\nUser Input: {user_input}\n\nPlease generate a full structured response, including Science Objectives."
    
    # Case 3: Neither context nor science objectives are provided
    elif science_objectives.strip():
        combined_input = f"User Input: {user_input}\nScience Objectives (User Provided): {science_objectives}\n\nPlease generate only the remaining sections as per the defined format."
    
    # Default: No relevant context or science objectives → Generate everything
    else:
        combined_input = f"User Input: {user_input}\n\nPlease generate a full structured response, including Science Objectives."
    
    response = client.chat.completions.create(
        model="gpt-4o",
        messages=[
            {"role": "system", "content": system_message},
            {"role": "user", "content": combined_input}
        ],
        max_tokens=max_tokens,
        temperature=temperature,
        top_p=top_p,
        frequency_penalty=frequency_penalty,
        presence_penalty=presence_penalty
    )

    response_only = response.choices[0].message.content.strip()

    # ADS References appended separately
    references_text = ""
    if references:
        references_text = "\n\nADS References:\n" + "\n".join(
            [f"- {title} {authors} (Bibcode: {bibcode}) {pub} {pubdate}" 
             for title, abstract, authors, bibcode, pub, pubdate in references])

    # Full response (for Gradio display)
    full_response = response_only + references_text

    # Return two clearly separated responses
    return full_response, response_only

def generate_data_insights(user_input, exoplanet_data, max_tokens=500, temperature=0.3):
    """
    Generate insights by passing the user's input along with the exoplanet data to GPT-4.
    """
    # Convert the dataframe to a readable format for GPT (e.g., CSV-style text)
    data_as_text = exoplanet_data.to_csv(index=False)  # CSV-style for better readability

    # Create a prompt with the user query and the data sample
    insights_prompt = (
        f"Analyze the following user query and provide relevant insights based on the provided exoplanet data.\n\n"
        f"User Query: {user_input}\n\n"
        f"Exoplanet Data:\n{data_as_text}\n\n"
        f"Please provide insights that are relevant to the user's query."
    )
    
    # Call GPT-4 to generate insights based on the data and user input
    response = client.chat.completions.create(
        model="gpt-4",
        messages=[
            {"role": "system", "content": "You are an expert in analyzing astronomical data and generating insights."},
            {"role": "user", "content": insights_prompt}
        ],
        max_tokens=max_tokens,
        temperature=temperature
    )
    
    # Extract and return GPT-4's insights
    data_insights = response.choices[0].message.content.strip()
    return data_insights


def export_to_word(response_content, subdomain_definition, science_goal, context, max_tokens, temperature, top_p, frequency_penalty, presence_penalty):
    doc = Document()
    
    # Add a title (optional, you can remove this if not needed)
    doc.add_heading('AI Generated SCDD', 0)

    # Insert the Subdomain Definition at the top
    doc.add_heading('Subdomain Definition:', level=1)
    doc.add_paragraph(subdomain_definition)

    # Insert the Science Goal at the top
    doc.add_heading('Science Goal:', level=1)
    doc.add_paragraph(science_goal)

    # Insert the User-defined Context
    doc.add_heading('User-defined Context:', level=1)
    doc.add_paragraph(context)

    # Insert Model Parameters
    doc.add_heading('Model Parameters:', level=1)
    doc.add_paragraph(f"Max Tokens: {max_tokens}")
    doc.add_paragraph(f"Temperature: {temperature}")
    doc.add_paragraph(f"Top-p: {top_p}")
    doc.add_paragraph(f"Frequency Penalty: {frequency_penalty}")
    doc.add_paragraph(f"Presence Penalty: {presence_penalty}")

    # Split the response into sections based on ### headings
    sections = response_content.split('### ')
    
    for section in sections:
        if section.strip():
            # Handle the "Observations Requirements Table" separately with proper formatting
            if 'Observations Requirements Table' in section:
                doc.add_heading('Observations Requirements Table', level=1)
                
                # Extract table lines
                table_lines = section.split('\n')[2:]  # Start after the heading line
                
                # Check if it's an actual table (split lines by '|' symbol)
                table_data = [line.split('|')[1:-1] for line in table_lines if '|' in line]
                
                if table_data:
                    # Add table to the document
                    table = doc.add_table(rows=len(table_data), cols=len(table_data[0]))
                    table.style = 'Table Grid'
                    for i, row in enumerate(table_data):
                        for j, cell_text in enumerate(row):
                            cell = table.cell(i, j)
                            cell.text = cell_text.strip()
                            # Apply text wrapping for each cell
                            cell._element.get_or_add_tcPr().append(parse_xml(r'<w:tcW w:w="2500" w:type="pct" ' + nsdecls('w') + '/>'))
                
                # Process any paragraphs that follow the table
                paragraph_after_table = '\n'.join([line for line in table_lines if '|' not in line and line.strip()])
                if paragraph_after_table:
                    doc.add_paragraph(paragraph_after_table.strip())
            
            # Handle the "ADS References" section
            elif section.startswith('ADS References'):
                doc.add_heading('ADS References', level=1)
                references = section.split('\n')[1:]  # Skip the heading
                for reference in references:
                    if reference.strip():
                        doc.add_paragraph(reference.strip())
            
            # Add all other sections as plain paragraphs
            else:
                doc.add_paragraph(section.strip())
    
    # Save the document to a temporary file
    temp_file = tempfile.NamedTemporaryFile(delete=False, suffix=".docx")
    doc.save(temp_file.name)
    
    return temp_file.name

def extract_table_from_response(gpt_response):
    # Split the response into lines
    lines = gpt_response.strip().split("\n")
    
    # Find where the table starts and ends (based on the presence of pipes `|` and at least 3 columns)
    table_lines = [line for line in lines if '|' in line and len(line.split('|')) > 3]
    
    # If no table is found, return None or an empty string
    if not table_lines:
        return None
    
    # Find the first and last index of the table lines
    first_table_index = lines.index(table_lines[0])
    last_table_index = lines.index(table_lines[-1])
    
    # Extract only the table part
    table_text = lines[first_table_index:last_table_index + 1]
    
    return table_text

def gpt_response_to_dataframe(gpt_response):
    # Extract the table text from the GPT response
    table_lines = extract_table_from_response(gpt_response)
    
    # If no table found, return an empty DataFrame
    if table_lines is None or len(table_lines) == 0:
        return pd.DataFrame()

    # Find the header and row separator (assume it's a line with dashes like |---|)
    try:
        # The separator line (contains dashes separating headers and rows)
        sep_line_index = next(i for i, line in enumerate(table_lines) if set(line.strip()) == {'|', '-'})
    except StopIteration:
        # If no separator line is found, return an empty DataFrame
        return pd.DataFrame()

    # Extract headers (the line before the separator) and rows (lines after the separator)
    headers = [h.strip() for h in table_lines[sep_line_index - 1].split('|')[1:-1]]
    
    # Extract rows (each line after the separator)
    rows = [
        [cell.strip() for cell in row.split('|')[1:-1]]
        for row in table_lines[sep_line_index + 1:]
    ]

    # Create DataFrame
    df = pd.DataFrame(rows, columns=headers)
    return df
    
def chatbot(user_input, science_objectives="", context="", subdomain="", max_tokens=150, temperature=0.7, top_p=0.9, frequency_penalty=0.5, presence_penalty=0.0):

    
    yield "🔄 Connecting with Pinecone...", None, None, None, None, None, None
    
    pc_index_name = "scdd-index"
    yield f"Using Pinecone index: **{index_name}**✅ ", None, None, None, None, None, None

    yield "🔎 Retrieving relevant context from Pinecone...", None, None, None, None, None, None
    # Retrieve relevant context using Pinecone
    relevant_context = retrieve_relevant_context(user_input, context, science_objectives)

    cleaned_context_list = [clean_retrieved_context(chunk) for chunk in relevant_context]
    

    yield "Context Retrieved successfully ✅ ", None, None, None, None, None, None, None

    keywords = extract_keywords_with_gpt(context)

    ads_query = " ".join(keywords)
    
    # Fetch NASA ADS references using the user context
    references = fetch_nasa_ads_references(ads_query)

    yield "ADS references retrieved... ✅ ", None, None, None, None, None, None, None
    

    yield "🔄 Generating structured response using GPT-4o...", None, None, None, None, None, None
    
    # Generate response from GPT-4
    full_response, response_only = generate_response(
        user_input=user_input,
        science_objectives=science_objectives,  
        relevant_context=relevant_context,
        references=references,
        max_tokens=max_tokens,
        temperature=temperature,
        top_p=top_p,
        frequency_penalty=frequency_penalty,
        presence_penalty=presence_penalty
    )

    # RAGAS Evaluation
    
    context_ragas = cleaned_context_list
    response_ragas = response_only
    query_ragas = user_input + context
    reference_ragas = "\n\n".join([f"{title}\n{abstract}" for title, abstract, _, _, _, _ in references])

    dataset = []

    dataset.append(
        {
            "user_input":query_ragas,
            "retrieved_contexts":context_ragas,
            "response":response_ragas,
            "reference":reference_ragas
        }
    )

    evaluation_dataset = EvaluationDataset.from_list(dataset)

    ragas_evaluation = evaluate(dataset=evaluation_dataset,metrics=[LLMContextRecall(), ContextRelevance(), Faithfulness(), ResponseRelevancy(), FactualCorrectness(coverage="high",atomicity="high")],llm=evaluator_llm, embeddings=embeddings)
    
    yield "Response generated successfully ✅ ", None, None, None, None, None, None
    
    # Append user-defined science objectives if provided
    if science_objectives.strip():
        full_response = f"### Science Objectives (User-Defined):\n\n{science_objectives}\n\n" + full_response

    # Export response to Word
    word_doc_path = export_to_word(
        full_response, subdomain, user_input, context, 
        max_tokens, temperature, top_p, frequency_penalty, presence_penalty
    )

    yield "Writing SCDD...Performing RAGAS Evaluation...", None, None, None, None, None, None
    
    # Fetch exoplanet data and generate insights
    exoplanet_data = fetch_exoplanet_data()
    data_insights = generate_data_insights(user_input, exoplanet_data)

    # Extract GPT-generated table into DataFrame
    extracted_table_df = gpt_response_to_dataframe(full_response)

    # Combine response and insights
    full_response = f"{full_response}\n\nEnd of Response"

    yield "SCDD produced successfully ✅", None, None, None, None, None, None

    iframe_html = """<iframe width=\"768\" height=\"432\" src=\"https://miro.com/app/live-embed/uXjVKuVTcF8=/?moveToViewport=-331,-462,5434,3063&embedId=710273023721\" frameborder=\"0\" scrolling=\"no\" allow=\"fullscreen; clipboard-read; clipboard-write\" allowfullscreen></iframe>"""
    mapify_button_html = """<a href=\"https://mapify.so/app/new\" target=\"_blank\"><button>Create Mind Map on Mapify</button></a>"""

    yield full_response, relevant_context, ragas_evaluation, extracted_table_df, word_doc_path, iframe_html, mapify_button_html

with gr.Blocks() as demo:
    gr.Markdown("# **ExosAI - NASA SMD PCRAG SCDD Generator [version-2.1]**")

    gr.Markdown("## **User Inputs**")
    user_input = gr.Textbox(lines=5, placeholder="Enter your Science Goal...", label="Science Goal")
    context = gr.Textbox(lines=10, placeholder="Enter Context Text...", label="Additional Context")
    subdomain = gr.Textbox(lines=2, placeholder="Define your Subdomain...", label="Subdomain Definition")

    science_objectives_button = gr.Button("User-defined Science Objectives [Optional]")
    science_objectives_input = gr.Textbox(lines=5, placeholder="Enter Science Objectives...", label="Science Objectives", visible=False)
    science_objectives_button.click(lambda: gr.update(visible=True), outputs=[science_objectives_input])

    gr.Markdown("### **Model Parameters**")
    max_tokens = gr.Slider(50, 2000, 150, step=10, label="Max Tokens")
    temperature = gr.Slider(0.0, 1.0, 0.7, step=0.1, label="Temperature")
    top_p = gr.Slider(0.0, 1.0, 0.9, step=0.1, label="Top-p")
    frequency_penalty = gr.Slider(0.0, 1.0, 0.5, step=0.1, label="Frequency Penalty")
    presence_penalty = gr.Slider(0.0, 1.0, 0.0, step=0.1, label="Presence Penalty")

    gr.Markdown("## **Model Outputs**")
    gr.Markdown("### **Accessing Pinecone vector database for context retrieval and generating response...**")
    full_response = gr.Textbox(label="ExosAI finds...")
    relevant_context = gr.Textbox(label="Retrieved Context...")
    ragas_evaluation = gr.Textbox(label="RAGAS Evaluation...")
    extracted_table_df = gr.Dataframe(label="SC Requirements Table")
    word_doc_path = gr.File(label="Download SCDD")
    iframe_html = gr.HTML(label="Miro")
    mapify_button_html = gr.HTML(label="Generate Mind Map on Mapify")

    with gr.Row():
        submit_button = gr.Button("Generate SCDD")
        clear_button = gr.Button("Reset")

    submit_button.click(chatbot, inputs=[user_input, science_objectives_input, context, subdomain, max_tokens, temperature, top_p, frequency_penalty, presence_penalty], outputs=[full_response, relevant_context, ragas_evaluation, extracted_table_df, word_doc_path, iframe_html, mapify_button_html],queue=True)

    clear_button.click(lambda: ("", "", "", "", 150, 0.7, 0.9, 0.5, 0.0, "", "", None, None, None, None, None), outputs=[user_input, science_objectives_input, context, subdomain, max_tokens, temperature, top_p, frequency_penalty, presence_penalty, full_response, relevant_context, ragas_evaluation, extracted_table_df, word_doc_path, iframe_html, mapify_button_html])

demo.launch(share=True)

